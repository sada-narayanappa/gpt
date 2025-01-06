#!/usr/bin/env python

import sys, os, logging,datetime, hashlib, pandas as pd
logger = logging.getLogger( "geoapp" )
from langchain_core.documents import Document
from mangorest.mango import webapi

# ------------------------------------------------------------------------------------------
def convertToDoc(file, out=None):
    import aspose.words as aw

    if not out:
        txt = open(file, "rb").read()
        h = hashlib.md5(txt)
        out = "/tmp" + str(datetime.datetime.now()).replace(" ","T") + ".docx"
        base = os.path.basename(file)
        out = f"/tmp/{base}-{(h.hexdigest())}.docx"

    if not os.path.exists(out):
        doc = aw.Document(file)
        doc.save(out)
    return out

SAMPLE_DOC = os.path.expanduser("~/Desktop/data/LLM/sample.docx")
SAMPLE_PDF = os.path.expanduser("~/Desktop/data/LLM/sample.pdf")

# ------------------------------------------------------------------------------------------
# Simply extarct text rfrom PDF file
def extractTextPDF(file):
    import pdfplumber

    #from pdfplumber.page import Page
    #from pdfplumber.pdf import PDF
    #from pdfplumber.table import Table
    #from pdfplumber.utils import intersects_bbox

    text = []
    #print(f"**** OPENINF FILE ===>  {f}")
    with pdfplumber.open(file) as doc:
        for page in doc.pages:
            #lines = page.extract_text_lines()
            #txt = "\n".join([l['text'] for l in lines])
            txt = page.extract_text_simple()
            text.append(txt)

    all="\n".join(text)
    return all

# ------------------------------------------------------------------------------------------
# Simply extarct text rfrom PDF file
def extractTextDOC(file):
    import docx
    document = docx.Document(file)
    txts=[]
    for p in document.paragraphs:
        txts.append(p.text)

    all = "\n".join(txts)
    return all

#-----------------------------------------------------------------------------------------    
@webapi("/gpt/extractText/")
def extractText(request=None, file=None, **kwargs):
    ret = f"Unknown file type {file}"

    if ( request and not file):
        for f in request.FILES.getlist('file'):
            content = f.read()
            #fileIO = io.BytesIO(content)
            file = f"/tmp/{str(f)}"
            with open(file, "wb") as f:
                f.write(content)


    if (file.endswith("doc") or file.endswith("docx") ):
        ret =  extractTextDOC(file)
    elif (file.endswith("pdf") ):
        ret = extractTextPDF(file)
    elif (file.endswith("xlsx") or file.endswith("xls")):
        df = pd.read_excel(file)
        ret = df.to_html()
    elif file.endswith("csv") :
        df = pd.read_csv(file)
        ret = df.to_html()
    elif file:
        ret = open(file, "rb").read()
    else:
        ret = ""
    
    return ret

# ---------------------------------------------------------------------------------------
def getChunks(file, chunk_size=2000, overlap=256 ):
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    
    txt= extractText( file= file)

    split = RecursiveCharacterTextSplitter(
        chunk_size= chunk_size,
        chunk_overlap=overlap,
        length_function=len,
        add_start_index=True,
    )
    docs = []
    for txt in split.split_text(txt):
        d = Document(page_content= txt,  metadata=dict(source= file))
        docs.append(d)

    return docs
